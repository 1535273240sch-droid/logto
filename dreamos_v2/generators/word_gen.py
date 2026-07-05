"""Word 生成器 — 生成 .docx 文件"""
from pathlib import Path
from . import BaseGenerator


class WordGenerator(BaseGenerator):
    supported_types = ["word"]

    async def generate(self, content: str, output_path: Path,
                       title: str = "", **kwargs) -> dict:
        self._ensure_dir(output_path)

        try:
            return await self._gen_docx(content, output_path, title)
        except ImportError:
            # 降级为 Markdown
            md_path = output_path.with_suffix(".md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(f"# {title}\n\n{content}")
            return {"success": True, "fallback": "markdown"}

    async def _gen_docx(self, content: str, output_path: Path,
                         title: str) -> dict:
        """使用 python-docx 生成 Word 文档"""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 标题
        title_para = doc.add_heading(title or "Document", level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 内容段落
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph("")
                continue

            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif len(line) > 2 and line[0].isdigit() and "." in line[:4]:
                doc.add_paragraph(line.split(".", 1)[1].strip(), style='List Number')
            else:
                para = doc.add_paragraph(line)
                para.style.font.size = Pt(11)

        doc.save(output_path)
        return {"success": True}
